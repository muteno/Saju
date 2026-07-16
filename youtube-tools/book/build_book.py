# -*- coding: utf-8 -*-
r"""분류 라벨(labels_final.json) + arcs(arcs_final.json) + 전문(fulltext) →
주제별 책 통합본 1개. refine-tools/assemble_book.py 를 자막용으로 옮긴 것.
- 17장(+C90 비명리 부록) 주제 분류, 이론장은 입문→심화 레벨 배열 + 장별 arc
- 전문은 인사말/구독유도만 가볍게 정제(내용 보존, 누락 금지)
출력: 바탕화면\유튜브 사주명리 대전 — 주제별 통합본.docx
사용: python build_book.py            (바탕화면 자동)
      python build_book.py "출력폴더"
"""
import json, re, sys
from pathlib import Path
from collections import defaultdict, Counter

HERE = Path(__file__).parent.resolve()
YT = HERE.parent

CHAPTERS = [
    ('C01', '제1장 명리 입문·기초', '음양오행 · 사주팔자 구조 · 만세력 · 절기'),
    ('C02', '제2장 천간과 지지', '10천간 · 12지지 · 지장간 · 물상'),
    ('C03', '제3장 합충형파해', '삼합·방합 · 충 · 형 · 파 · 해'),
    ('C04', '제4장 십신·육친론', '비겁 · 식상 · 재성 · 관성 · 인성'),
    ('C05', '제5장 일주론', '60일주 · 일간별 특성'),
    ('C06', '제6장 신살론', '귀인 · 도화 · 역마 · 화개 · 괴강 등'),
    ('C07', '제7장 십이운성', '장생에서 양까지 · 왕상휴수'),
    ('C08', '제8장 원국 분석', '신강신약 · 용신 · 격국 · 조후'),
    ('C09', '제9장 운 해석 방법론', '대운·세운·월운 보는 법과 원리'),
    ('C10', '제10장 연운 — 신년운세', '연도별·일주별 한 해 운세'),
    ('C11', '제11장 월운·일운', '월간운세 · 이달의 운세'),
    ('C12', '제12장 실전 풀이·상담 사례', '실제 사주 풀이 · 상담 사례'),
    ('C13', '제13장 명리 칼럼·에세이', '사주와 과학 · 역사 · 생활 속 명리'),
    ('C14', '제14장 고전 강독', '자평류 고전 · 추명가'),
    ('C15', '제15장 재미로 읽는 사주', '유명인 · 로또 · 영화/TV · 신변잡기'),
    ('C16', '제16장 기타 역학', '풍수 · 타로 · 관상 · 수상 · 구성학 등'),
    ('C17', '제17장 부록 — 이용 안내', '사용설명서 · 수강/서비스 · 공지'),
    ('C90', '제18장 부록 — 비명리 영상', '뮤직비디오 · 뉴스 · 스포츠 등(참고 보관)'),
]
CH_TITLE = {c: (t, s) for c, t, s in CHAPTERS}
LEVELED = {'C01', 'C02', 'C03', 'C04', 'C05', 'C06', 'C07', 'C08', 'C09', 'C12', 'C13', 'C14'}
LEVEL_NAME = {1: '입문', 2: '초급', 3: '중급', 4: '심화'}
# 강의 채널 우선순위(장 내부, 같은 레벨에서 계통 유지)
CH_PRI = {'초코서당(초명)': 0, '도화도르- 사주팔자 쉽게 풀어주는 남자': 1,
          '마인드 명리': 2, '포스텔러 FORCETELLER': 3, '길 인간학연구소': 4}

GREET_LEAD = re.compile(
    r'^\s*(네[.,~\s]*)?(자[,\s])?(여러분[,\s]*)?'
    r'(안녕하세요[^.!?]{0,45}[.!?]\s*)+(반갑습니다[^.!?]{0,25}[.!?]\s*)?')
OUTRO = re.compile(r'(구독|좋아요|알림\s?설정|다음\s?(시간|영상|편)에|시청해\s?주셔서|'
                   r'다음에\s?(또\s?)?(뵙|만나|찾아)|영상은\s?여기까지|이상\s?\S{0,6}이었습니다)')


def clean_paras(paras):
    paras = [p.strip() for p in paras if p.strip()]
    if paras:
        paras[0] = GREET_LEAD.sub('', paras[0]).strip()
        if not paras[0]:
            paras.pop(0)
    while paras and OUTRO.search(paras[-1]) and len(paras[-1]) < 180:
        paras.pop()
    return paras


def find_desktop():
    if len(sys.argv) > 1:
        return Path(sys.argv[1])
    home = Path.home()
    for name in ("Desktop", "바탕 화면", "바탕화면", "OneDrive/Desktop"):
        p = home / name
        if p.exists():
            return p
    return home / "Desktop"


YEAR_RX = re.compile(r'(20\d\d)')
MONTH_RX = re.compile(r'(\d{1,2})\s*월')


def order_key(cat, u, level):
    ch = CH_PRI.get(u['channel'], 50)
    date = u.get('date') or '99999999'
    if cat in LEVELED:
        return (level or 9, ch, date)
    if cat == 'C10':
        y = YEAR_RX.search(u['title']) or YEAR_RX.search(date)
        return (int(y.group(1)) if y else 9999, ch, date)
    if cat == 'C11':
        y = YEAR_RX.search(u['title'])
        m = MONTH_RX.search(u['title'])
        return (int(y.group(1)) if y else 9999, int(m.group(1)) if m else 99, date)
    return (ch, date)


def main():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    units = {u['vid']: u for u in json.load(open(HERE / 'yt_units.json', encoding='utf-8'))}
    labels = json.load(open(HERE / 'labels_final.json', encoding='utf-8'))
    arcs = json.load(open(HERE / 'arcs_final.json', encoding='utf-8')) if (HERE / 'arcs_final.json').exists() else {}

    lab = {l['vid']: l for l in labels}
    # 커버리지 assert (누락 금지)
    missing = [v for v in units if v not in lab]
    assert not missing, f'라벨 누락 {len(missing)}건: {missing[:8]}'
    bad = {v: l['cat'] for v, l in lab.items() if l['cat'] not in CH_TITLE}
    assert not bad, f'잘못된 카테고리: {list(bad.items())[:5]}'

    by_ch = defaultdict(list)
    for v, l in lab.items():
        by_ch[l['cat']].append(v)
    for c in by_ch:
        by_ch[c].sort(key=lambda v: order_key(c, units[v], lab[v].get('level', 0)))

    GREEN = RGBColor(0x14, 0x53, 0x2D)
    GREEN2 = RGBColor(0x16, 0x65, 0x34)
    GRAY = RGBColor(0x66, 0x66, 0x66)
    GRAY2 = RGBColor(0x88, 0x88, 0x88)
    LINK = RGBColor(0x1a, 0x0d, 0xab)

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = '맑은 고딕'
    normal.font.size = Pt(10.5)

    def run(par, text, size=None, color=None, bold=False, italic=False):
        r = par.add_run(text)
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        r.bold = bold; r.italic = italic
        return r

    present = [c for c, _, _ in CHAPTERS if by_ch.get(c)]

    # ── 표지 ──
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, '유튜브 사주명리 대전', 26, GREEN, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, f'주제별 통합본 · 전 {len(present)}장 · 총 {len(units):,}편', 13, GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, '유튜브 사주명리 강의 자막을 주제별로 재분류하고 입문→심화 순으로 엮음', 10.5, GRAY2, italic=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, '이론 장은 입문·초급·중급·심화 구분선으로 배열 · 각 자막은 원문 주소 표기', 9, GRAY2, italic=True)

    # ── 목차 ──
    doc.add_page_break()
    h = doc.add_paragraph(); run(h, '목차', 18, GREEN, bold=True)
    for c, t, s in CHAPTERS:
        n = len(by_ch.get(c, []))
        if not n:
            continue
        p = doc.add_paragraph()
        run(p, t, 11, GREEN2, bold=True)
        run(p, f'  — {s} · {n}편', 9, GRAY2)

    # ── 본문 ──
    seq = 0
    for c, t, s in CHAPTERS:
        entries = by_ch.get(c, [])
        if not entries:
            continue
        doc.add_page_break()
        p = doc.add_paragraph(); run(p, t, 22, GREEN, bold=True)
        p = doc.add_paragraph(); run(p, f'{s} · {len(entries)}편', 12, GRAY)
        arc = arcs.get(c)
        if arc:
            p = doc.add_paragraph(); run(p, f'진행 방향 — {arc}', 10, RGBColor(0x77, 0x77, 0x77), italic=True)
        cur_lv = None
        for v in entries:
            u = units[v]
            if c in LEVELED:
                lv = lab[v].get('level') or 0
                if lv and lv != cur_lv:
                    pl = doc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run(pl, f'· {LEVEL_NAME.get(lv, str(lv))} ·', 15, GREEN2, bold=True)
                    cur_lv = lv
            seq += 1
            ph = doc.add_paragraph(); run(ph, f'{seq:04d}. {u["title"]}', 14, GREEN2, bold=True)
            sub = lab[v].get('subtopic', '')
            meta = doc.add_paragraph()
            run(meta, f'채널: {u["channel"]}    업로드: {u["date"]}    길이: {u["dur"]}'
                + (f'    · {sub}' if sub else ''), 9, GRAY)
            pl = doc.add_paragraph(); run(pl, u['url'], 9, LINK)
            ft = HERE / 'fulltext' / f'{v}.txt'
            paras = ft.read_text(encoding='utf-8').split('\n\n') if ft.exists() else []
            for para in clean_paras(paras):
                doc.add_paragraph(para)

    out = find_desktop() / '유튜브 사주명리 대전 — 주제별 통합본.docx'
    doc.save(str(out))

    dist = Counter(lab[v]['cat'] for v in units)
    print('[분포]')
    for c, t, s in CHAPTERS:
        if dist.get(c):
            print(f'  {c} {t}: {dist[c]}편')
    print(f'[done] {out}  ·  {out.stat().st_size/1024/1024:.2f} MB  ·  {seq}편 수록')
    assert seq == len(units), f'수록 {seq} != 전체 {len(units)}'


if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    main()
