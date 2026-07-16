# -*- coding: utf-8 -*-
"""이미 받아둔 raw/ 자막(ko.json3)+메타(info.json)를 바탕화면 산출물로 조립.
네트워크 호출 없음. to_text.py 의 json3->문단 로직을 그대로 재사용한다.

산출물 (바탕화면/유튜브_자막_모음/):
  - 유튜브_자막_모음.docx : 영상별 제목+메타+자막을 담은 단일 워드
  - 결과요약.csv          : 전체 영상 성공/실패 표 (utf-8-sig, 엑셀 호환)
  - 텍스트/NNNN_제목.txt  : 자막 있는 영상 1개당 파일 1개

사용:
  python build_desktop.py                # 바탕화면 자동
  python build_desktop.py "D:\\원하는\\폴더"  # 대상 지정
"""
import os, re, sys, csv, json
from collections import Counter
from pathlib import Path

HERE = Path(__file__).parent.resolve()
sys.path.insert(0, str(HERE))
from to_text import json3_to_paragraphs  # 검증된 변환 로직 재사용

RAW = HERE / "raw"
SUB_VARIANTS = ("ko", "ko-orig", "ko-KR")


def find_desktop():
    if len(sys.argv) > 1:
        return Path(sys.argv[1])
    home = Path.home()
    for name in ("Desktop", "바탕 화면", "바탕화면", "OneDrive/Desktop", "OneDrive/바탕 화면"):
        p = home / name
        if p.exists():
            return p
    return home / "Desktop"


def sanitize(name, limit=80):
    name = re.sub(r'[\\/:*?"<>|\r\n\t]', " ", name or "")
    name = re.sub(r"\s+", " ", name).strip(" .")
    return (name[:limit].strip() or "무제")


def load_json(p):
    try:
        return json.load(open(p, encoding="utf-8"))
    except Exception:
        return {}


def sub_path(vid):
    for s in SUB_VARIANTS:
        p = RAW / f"{vid}.{s}.json3"
        if p.exists():
            return p
    return None


def error_reasons():
    """fetch.log / fetch_missing.log 에서 영상별 실패 사유를 추출."""
    reasons = {}
    for name in ("fetch.log", "fetch_missing.log"):
        p = HERE / name
        if not p.exists():
            continue
        for line in p.open(encoding="utf-8", errors="replace"):
            low = line.lower()
            if "error" not in low and "no-ko" not in low:
                continue
            m = re.search(r"\]\s*([A-Za-z0-9_-]{11}):", line) or \
                re.search(r"no-ko\s+([A-Za-z0-9_-]{11})", line)
            if not m:
                continue
            vid = m.group(1)
            if "member" in low:
                reasons[vid] = "멤버십 전용(접근 불가)"
            elif "confirm your age" in low or "age-restricted" in low or "inappropriate" in low:
                reasons.setdefault(vid, "연령 제한(로그인 필요)")
            elif "confirm you're not a bot" in low or "sign in to confirm you" in low:
                reasons.setdefault(vid, "봇 확인 차단(재시도 대상)")
    return reasons


def main():
    master = load_json(HERE / "master.json")
    all_ids = master.get("all_ids") or []
    if not all_ids:
        print("master.json 에서 all_ids 를 찾지 못했습니다.")
        sys.exit(1)

    pl_of = {}
    for pid, pl in (master.get("playlists") or {}).items():
        for vid in pl.get("ids", []):
            pl_of[vid] = pl.get("title") or pid

    reasons = error_reasons()

    out = find_desktop() / "유튜브_자막_모음"
    txt_dir = out / "텍스트"
    txt_dir.mkdir(parents=True, exist_ok=True)
    for old in txt_dir.glob("*.txt"):
        old.unlink()

    rows, docx_items = [], []
    seq = 0
    ok = 0
    fails = Counter()

    for vid in all_ids:
        info_p = RAW / f"{vid}.info.json"
        sub_p = sub_path(vid)
        i = load_json(info_p) if info_p.exists() else {}
        title = i.get("title") or ""
        channel = i.get("channel") or i.get("uploader") or ""
        date = i.get("upload_date") or ""
        dur = i.get("duration_string") or ""
        url = i.get("webpage_url") or f"https://www.youtube.com/watch?v={vid}"
        playlist = pl_of.get(vid, "")

        paras, nchars, nparas = [], 0, 0
        status, reason = "실패", ""

        if sub_p is not None:
            paras = json3_to_paragraphs(str(sub_p))
            nparas = len(paras)
            nchars = sum(len(p) for p in paras)
            if nchars > 0:
                status, reason = "성공", ""
                ok += 1
            else:
                reason = "빈 자막(내용 없음)"
        elif vid in reasons:
            reason = reasons[vid]
        elif info_p.exists():
            reason = "자막 없음(한국어 자막 미제공)"
        else:
            reason = "미수집(재수집 대상)"

        if status == "실패":
            fails[reason] += 1

        if status == "성공":
            seq += 1
            fname = f"{seq:04d}_{sanitize(title)}.txt"
            body = (
                f"제목: {title}\n"
                f"채널: {channel}\n"
                f"업로드: {date}    길이: {dur}\n"
                f"주소: {url}\n"
                f"영상ID: {vid}"
                + (f"\n재생목록: {playlist}" if playlist else "")
                + "\n" + "=" * 60 + "\n\n"
                + "\n\n".join(paras) + "\n"
            )
            (txt_dir / fname).write_text(body, encoding="utf-8")
            docx_items.append((seq, {
                "title": title, "channel": channel, "date": date,
                "dur": dur, "url": url, "vid": vid, "playlist": playlist,
            }, paras))

        rows.append({
            "번호": seq if status == "성공" else "",
            "영상ID": vid, "상태": status, "사유": reason,
            "채널": channel, "제목": title, "업로드일": date,
            "길이": dur, "재생목록": playlist,
            "문단수": nparas, "글자수": nchars, "주소": url,
        })

    # --- CSV ---
    csv_p = out / "결과요약.csv"
    cols = ["번호", "영상ID", "상태", "사유", "채널", "제목",
            "업로드일", "길이", "재생목록", "문단수", "글자수", "주소"]
    with open(csv_p, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=cols)
        w.writeheader()
        w.writerows(rows)

    # --- DOCX ---
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10.5)

    total = len(all_ids)
    h = doc.add_heading("유튜브 자막 모음", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run(
        f"자막 수록 영상 {len(docx_items)}편 · 전체 대상 {total}편\n"
        f"영상별 제목 · 채널 · 주소 · 전체 자막 수록"
    ).italic = True
    doc.add_page_break()

    for seq_i, m, paras in docx_items:
        doc.add_heading(f"{seq_i:04d}. {m['title']}", level=1)
        meta = doc.add_paragraph()
        r = meta.add_run(
            f"채널: {m['channel']}    업로드: {m['date']}    길이: {m['dur']}"
            + (f"    재생목록: {m['playlist']}" if m["playlist"] else "")
        )
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        link = doc.add_paragraph()
        lr = link.add_run(m["url"])
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0x1a, 0x0d, 0xab)
        for p in paras:
            doc.add_paragraph(p)
        doc.add_page_break()

    docx_p = out / "유튜브_자막_모음.docx"
    doc.save(str(docx_p))

    # --- 요약 ---
    print("=" * 60)
    print(f"완료: 성공 {ok} / 전체 {total}")
    print(f"  - txt 생성: {ok}편, 워드 수록: {len(docx_items)}편")
    print(f"실패 {sum(fails.values())}편 내역:")
    for reason, n in fails.most_common():
        print(f"  - {reason}: {n}")
    print("-" * 60)
    print(f"출력 폴더 : {out}")
    print(f"  워드     : {docx_p.name}  ({docx_p.stat().st_size/1024/1024:.2f} MB)")
    print(f"  요약 CSV : {csv_p.name}")
    print(f"  텍스트   : 텍스트\\ ({ok}개 파일)")


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    main()
