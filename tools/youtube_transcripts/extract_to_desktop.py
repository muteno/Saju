#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
유튜브 자막(대화 글자) 전체 추출 → 바탕화면 폴더에 영상별 저장 + 워드 1개
======================================================================
■ 이 파일 하나면 됩니다 (독립 실행 · 리스트 내장 · 레포 없어도 동작).

무엇을 하나:
  1) 아래 VIDEO_IDS / PLAYLISTS(영상 목록)를 하나씩 돌며 자막을 통째로 가져온다.
  2) 바탕화면(Desktop)에 '유튜브_자막_모음' 폴더를 자동으로 만든다.
  3) 영상 1개당 텍스트 파일 1개로 저장한다  →  바탕화면/유튜브_자막_모음/텍스트/0001_제목.txt
  4) 전부 묶은 워드 파일 1개를 만든다        →  바탕화면/유튜브_자막_모음/유튜브_자막_모음.docx
  5) 성공/실패 요약표(CSV)도 남긴다.

왜 '내 컴퓨터'에서 돌리나:
  유튜브는 클라우드 서버(데이터센터) IP의 자동 자막 요청을 봇으로 보고 막는다.
  (확인됨: 서버에서 요청하면 google.com/sorry 차단 페이지로 튕김.)
  집·회사 일반 인터넷에서는 대부분 정상 동작한다. 그래서 이 파일을 내 컴퓨터에서 실행한다.

실행 방법 (터미널/명령프롬프트에서):
  1) 필요한 부품 설치 (한 번만):
       pip install youtube-transcript-api python-docx yt-dlp
  2) 실행:
       python extract_to_desktop.py

  ─ 특정 폴더에 넣고 싶으면:  python extract_to_desktop.py "C:\\내가\\원하는\\폴더"
"""

import os
import re
import sys
import time
import json
import html
import urllib.request
from pathlib import Path

# ════════════════════════════════════════════════════════════════════════
#  영상 목록 (사용자가 보낸 주소에서 중복 제거함) — 아래 두 리스트가 곧 작업 대상
# ════════════════════════════════════════════════════════════════════════
VIDEO_IDS = [
    "7wLWFYMflaw", "hdUj1lq2jXc", "yUcasHEzOzQ", "BfPwayJ1n0A", "1ZEmQmYM8bU", "dTzS50z6s7g",
    "KwivNLfoGbE", "rNqihUidxvI", "W0XO1ylzN7o", "vMdsdhZqQ_0", "YNrxFQj_FPY", "ahBYzr2eZn0",
    "qRHyAyjFZLs", "8Pf3hGzVprE", "kDvPEALXeLU", "1QFaCE4FucA", "uIfCv_XngN8", "KiV4DcjtMpA",
    "bFfowUX0rjY", "Al8o55wAqy4", "A8rWK0luiWQ", "jiZdHvzKfv4", "QlCGgC-Ocqk", "KLDPWKsjhMo",
    "mAk12tAOSJY", "gbDyXTTHTxs", "z_MqlchyWPg", "pRQRW6tRcDw", "y8Vgcp4H4I4", "EKzoXPyjAac",
    "UAQGqOd6wRw", "sC3-iYHCfaU", "DmbuIwudvEk", "aQGmn9bzHss", "mgiNlvICLJk", "uAlRQiKyBKA",
    "-YSebkN6jsA", "rF5UexsQCj8", "yWI3zFxW6g0", "aQR2Um81E_4", "WzDcx1eYpLI", "_JBE5OOnqBo",
    "EOYvyD83unM", "GsJk0jrmPqQ", "RxSNnEscbEY", "wswodcBqsQo", "uvxYs4IH10Q", "0mHtSAb1Mpo",
    "D7Mt2UrTYGU", "FELhO8z7F94", "zo73W1cQhik", "MoWx3evGliQ", "vafUlOzUE9E", "l2LLa4IkJC0",
    "dHdc0RWTGls", "vQU1XtmYmJk", "X7NvVDd2exQ", "6vO-fHGeLdQ", "uzlDfStJJk0", "1WAy2D5nVCc",
    "E1FqaMsD3cw", "7ilwh0p5lpY", "qW1zhy8sTsg", "gfwka248_Vg", "vrZY6CX5Dy0", "eBt3vuk_KpU",
    "nTkTpmOx-Q8", "ffTo4sCnOHg", "Ibz48WPopYM", "IeAl9HkHxhg", "cX7ZQ-MoOFk", "WiY1ku_AqNg",
    "N8oqrpRS8uw", "O_boVjTSnkA", "AcTC59jt8qs", "jhI9lkSGFKw", "_dztv4Uw95I", "o4IkXat87cs",
    "sdGUxQ-eiI4", "g_pf1FHQGVs", "mEpzUdYQEaQ", "LTnuvjiwvLY", "a3qufG3oEug", "J6yMBhXevK8",
    "5nHIVlZcpdA", "wWFGPqEbinc", "Q69NMW8ceEc", "nNuYJ3WYNIU", "rXNlPbwUt3Q", "MKn_My9P2_s",
    "obnysj4TEv8", "1mRtMD6TXQg", "Hdn_wEMauGQ", "nygIAOJA5nw", "GY_jTMvvxL0", "UqPE99sviz0",
    "VSuNGpYP_RM", "UGiCEV6yYug", "GvtNBLn6_Ik", "OQoGuu4D_gI", "adbzlafaesU", "XkkEM40VwFo",
    "cz62--Of5dc", "Apxi-9FfI3M", "jY6TCArL3hY", "I1ZESmsJ8l0", "yUZosfgYJdQ", "FJMxMJkHhnA",
    "qLIgOoA-hZ0", "4UOFq6XHt60", "74G_oLUOadU", "P6Ai83qwPJo", "onBX2HMx0d4", "t8LY1Y9Nntg",
    "2GJfWMYCWY0", "abskiSwixec", "7YJ_C9Yuoz4", "83C3TZ4Zm_o", "ao7AymCB6uc", "TXPK9N0AEcY",
    "6D0CxKlD9Pw", "DnT1933rvrE", "sH22ypRpjf8", "CQAoatYfCQk", "aAZwwo-3GcM", "TI3PIiQyZZI",
    "uvzD3CWbEAo", "PtZbfqDgIvY", "dZs_cLHfnNA", "JJ84ZyVoH0o", "4rKeGsHa0lE", "2NVrg-umX84",
    "RoDm_p9lZRE", "_hjMxDrak-U", "59ZSlonVGOo", "KQRG9KvLjSk", "ODZBXHGOt84", "raSgQTSzoAo",
    "TsXTxYrFzFU", "HNAAyIm-1Qo", "nbOXfQSORi8", "LO0NEY67RAs", "1wKAbEVTjv4", "0mZbb67cJ-U",
    "wwpzUWYVlIE", "y0UjNHtNpvw", "62e69Bz7-ww", "4lGGG-y42EY", "z89FlpMnqp8", "7e86ERJmt5o",
    "ekBG8hc32Lo", "ImJ8ARLH6cg", "zFgTCU1HF1w", "9XttLI0oH0I", "ptrGHOngF_E", "VtI5Og0ktok",
    "OMtnPc_jaGo", "3RwqFxM2Ews", "cewJ4bk6_Hc", "hLvWy2b857I", "jLYxdnzdNfE", "lBKEpXwtpkQ",
    "n_Iqp5bhvtY", "yf_wpcH3Diw", "ADK7WqK4d1Y", "HFNZmoinSmQ", "79K-rIkJSNU", "9w_nPKVm8Bk",
    "97xRBoAAij0", "BoP0bxW3XAo", "4-TZbMZVpdQ", "uzZjlpOzAX4", "gDDyZkXcfpI", "GjTshpZWY5E",
    "Aks01uMljX4", "z2qCI8DZa28", "U2lkItDJ4kE", "6ljooVVZpps", "C1_8B9XOAeg", "fblqpRXNgUk",
    "aEoXEt1Dc0I", "PVaGLhG3cR4", "6mM3TmvcLKU", "FBGU4nJieME", "HbfDN7ixAz4", "roTt1E0P4hI",
    "Rc24FhNWO48", "5N8Nmdzbmv0", "iGUblr7Ts38", "JchNg3hnHtw", "5FicObGJHlY", "7UPOzS4oJms",
    "TSyBimyglxQ", "S--sJwiLI6M", "S8WRMGzni0k", "tIjb3ZaKh7c", "Akl_furFk7k", "Q7TbralZp5w",
    "oE_5Lruwe30", "L1FB3EWqkrw", "6UZ1yJCucA0", "nBYU5N6MIVM", "A7MwVr95OvA", "d79uykAeN8E",
    "sIdlhPO2Ado", "sGZhl-HQTKA", "iutB4e0Aq5g", "m7cf9FYk0Jw", "MtP9MSn9hl4", "jxoxkl1WZHM",
    "kKsivrgoyDw", "3VTkBuxU4yk", "Os_heh8vPfs", "phuiiNCxRMg", "M2WTUoy4y6E", "WAUSsfU00L0",
    "7-2n8booJYY", "NbuyCF9jRn4", "PtxDCBcBNhM", "fCLH1u1mxtk", "0lb2zWOuMII", "6ZUIwj3FgUY",
    "ekr2nIex040", "nFYwcndNuOY", "2iK3ccCsI6s", "UOxkGD8qRB4", "2VAKcz9W4WM", "EAAnf5KRcXY",
    "kW816TER06A", "03bEv3v2bSk", "rrMyqH18ddA", "ItTc8_5Vh44", "olGpmlwxYac", "agW3gtY6Huk",
    "UqPsWntdd1U", "u_4OMBtthtY", "Esq15Nk5lnU", "fJq1vT_9BbU", "7yMisDZugYQ", "DdKB1jQALPw",
    "jtjVORGeGFg", "2y2Z06hVmWE", "Drc0SgoedMA", "Rn89mUndYbo", "qORaYudQ7Zc", "4Zp5j1fN9Xw",
    "ICNKal4BIFI", "eKfJ3y41RDU", "K7bPpghHj3Q", "r5oAQkNvMGk", "KOO_2eh1VJQ", "DU-Ilbnjoz4",
    "IylGkAcEW2o", "MGTQLfn3kvE", "pBUlV7rZp_c", "C3GouGa0noM", "B1ShLiq3EVc", "bMhDJ0S0OBA",
    "DYgE3SGPEqk", "iTJSbJtS8MU", "F0B7HDiY-10", "jOTfBlKSQYY", "CgCVZdcKcqY", "38xYeot-ciM",
    "g36q0ZLvygQ", "1Lmy7qwmSMc", "GU-26HelNgk", "4YuO8HXDvXo", "tS9jTiz2qSA", "6IFeGqSqrKs",
    "TPjq_IXUpkE", "x2eJlMatzK4", "Z1Axn8M5Fcs", "8AVb4nPI_5Y", "ZZ7kljLImck", "gR9gDyLbYNo",
    "5spDxl-yzUQ", "wdH5_I7UiHA", "3hEuDhAnVNc", "SZiwpL62to8", "4N15045PHEA", "Wk3Vxt3knnI",
    "ewJBkYws3AA", "WdiSosDz4ss", "9JFi7MmjtGA", "sVTy_wmn5SU", "Rrf8uQFvICE", "PGLx4V680J8",
    "NYT3Opelerg", "9qkpcLK422o", "JyvvqnUaMgY", "n6B5gQXlB-0", "OpqB1GJbzms", "6oG3oaZ4HQg",
    "rqA5vHcZ8jY", "zoiuqxsRx8U", "zy9VZ9ZMmJc", "iW9cycjnSvc", "iQbLJJ8tM6Q", "VYx4uU8qqRU",
    "ZICbCFTWeTM", "vpg_PnONB_s", "pLHom8xqMvA", "chAlIgS5PwA", "EmeW6li6bbo", "NIR9CQTEB_g",
    "h3rmoTooBnc", "vHAGhAAolMg", "DRzVaV8cyDk", "dB109Z-vza0", "aPOzOf44LEc", "3CgIhN7Aupw",
    "So-8rVI3ZKk", "mgFyfe_nW58", "wxTMCgxJpsY", "AtKe-X-lWWM", "IP1jk3ODH_Y", "8LaeFr7vXiM",
    "bWNbbV-1ssU", "nWtC8Mk1l0w", "a1IuidFKfLk",
]
PLAYLISTS = [
    "PLe8KqGCNSq6exRhWnyewUwVDCMAIYMsWJ", "PLe8KqGCNSq6e2KLvb62cEV_uWBxHoUV2a", "PLe8KqGCNSq6cYK3jyj4QU8Zu4UiuQxf9t", "PLe8KqGCNSq6cfWqDGp-KCcerbi82ZHMsi",
]

# ─────────────────────────── 설정 (원하면 여기만 조정) ───────────────────────────
FOLDER_NAME = "유튜브_자막_모음"                    # 바탕화면에 만들 폴더 이름
LANG_PREF = ["ko", "ko-KR", "en", "en-US"]         # 자막 우선순위: 한국어 먼저, 없으면 영어
SLEEP_SEC = 1.0                                     # 영상 사이 쉬는 시간(초) — 차단 회피
RETRY = 2                                           # 실패 시 재시도 횟수
TIMESTAMPS = False                                  # True면 각 줄 앞에 [분:초] 시간 표시
EXPAND_PLAYLISTS = True                             # 재생목록을 개별 영상으로 펼칠지 (yt-dlp 필요)
# ──────────────────────────────────────────────────────────────────────────────


def log(msg):
    print(msg, flush=True)


def find_desktop():
    """윈도우/맥/리눅스 + OneDrive + 한글 '바탕화면'까지 바탕화면 폴더를 자동 탐지."""
    home = Path.home()
    candidates = [
        home / "Desktop",
        home / "바탕화면",
        home / "바탕 화면",
        home / "OneDrive" / "Desktop",
        home / "OneDrive" / "바탕 화면",
        home / "OneDrive" / "바탕화면",
    ]
    up = os.environ.get("USERPROFILE")
    if up:
        candidates.insert(0, Path(up) / "Desktop")
        candidates.append(Path(up) / "OneDrive" / "Desktop")
    for c in candidates:
        try:
            if c.is_dir():
                return c
        except OSError:
            continue
    # 못 찾으면 홈 폴더 밑에 Desktop을 만들어 사용
    fallback = home / "Desktop"
    fallback.mkdir(parents=True, exist_ok=True)
    return fallback


def expand_playlist(playlist_id):
    """yt-dlp로 재생목록 안의 영상 ID들을 펼친다. yt-dlp 없으면 빈 목록."""
    try:
        import yt_dlp
    except ImportError:
        log(f"  [건너뜀] yt-dlp 없음 → 재생목록 못 펼침: {playlist_id}  (pip install yt-dlp 하면 됨)")
        return []
    ids = []
    opts = {"quiet": True, "no_warnings": True, "extract_flat": True, "skip_download": True}
    try:
        with yt_dlp.YoutubeDL(opts) as ydl:
            info = ydl.extract_info(f"https://www.youtube.com/playlist?list={playlist_id}", download=False)
            for e in info.get("entries", []) or []:
                vid = e.get("id")
                if vid and re.fullmatch(r"[A-Za-z0-9_-]{11}", vid):
                    ids.append(vid)
        log(f"  재생목록 {playlist_id}: {len(ids)}개 영상 펼침")
    except Exception as e:
        log(f"  [실패] 재생목록 {playlist_id}: {type(e).__name__} {str(e)[:120]}")
    return ids


def get_title(video_id):
    """oEmbed로 제목/채널을 가볍게 가져온다 (API 키 불필요). 실패하면 ID로 대체."""
    url = f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={video_id}&format=json"
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        data = json.loads(urllib.request.urlopen(req, timeout=15).read().decode("utf-8"))
        return data.get("title") or video_id, data.get("author_name") or ""
    except Exception:
        return video_id, ""


def _fmt_ts(sec):
    sec = int(sec)
    return f"[{sec // 60:02d}:{sec % 60:02d}] "


def fetch_transcript(video_id):
    """자막을 가져온다. 반환: (텍스트, 언어, 방법) / 실패 시 (None, None, 사유)."""
    # ── 1순위: youtube-transcript-api (신/구 버전 모두 대응) ──
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        snippets = None
        used_lang = None
        try:  # 신 인터페이스 (v1.x)
            api = YouTubeTranscriptApi()
            fetched = api.fetch(video_id, languages=LANG_PREF)
            snippets = [(s.text, s.start) for s in fetched]
            used_lang = getattr(fetched, "language_code", "?")
        except TypeError:  # 구 인터페이스 (v0.6.x)
            data = YouTubeTranscriptApi.get_transcript(video_id, languages=LANG_PREF)
            snippets = [(d["text"], d.get("start", 0)) for d in data]
            used_lang = LANG_PREF[0]
        except Exception:  # 우선언어 없음 → 아무 자막이나 (한국어로 번역 시도)
            try:
                api = YouTubeTranscriptApi()
                tlist = api.list(video_id)
                tr = None
                try:
                    tr = tlist.find_transcript(LANG_PREF)
                except Exception:
                    for t in tlist:
                        tr = t
                        break
                if tr is not None:
                    try:
                        if tr.language_code not in ("ko", "ko-KR") and tr.is_translatable:
                            tr = tr.translate("ko")
                    except Exception:
                        pass
                    f2 = tr.fetch()
                    snippets = [(s.text, s.start) for s in f2]
                    used_lang = getattr(tr, "language_code", "?")
            except Exception:
                snippets = None

        if snippets:
            if TIMESTAMPS:
                text = "\n".join(_fmt_ts(st) + html.unescape(t).replace("\n", " ") for t, st in snippets)
            else:
                text = "\n".join(html.unescape(t).replace("\n", " ") for t, _ in snippets)
            return text.strip(), used_lang, "youtube-transcript-api"
    except ImportError:
        log("  [경고] youtube-transcript-api 미설치 → pip install youtube-transcript-api")
    except Exception:
        pass

    # ── 2순위: yt-dlp 자막 백업 ──
    try:
        text, lang = fetch_transcript_ytdlp(video_id)
        if text:
            return text, lang, "yt-dlp"
    except Exception:
        pass

    return None, None, "자막 없음/차단"


def fetch_transcript_ytdlp(video_id):
    """yt-dlp로 자막 트랙 URL을 얻어 timedtext(json3)를 직접 받아 텍스트로 변환."""
    import yt_dlp
    opts = {"skip_download": True, "quiet": True, "no_warnings": True,
            "writesubtitles": True, "writeautomaticsub": True}
    with yt_dlp.YoutubeDL(opts) as ydl:
        info = ydl.extract_info(f"https://www.youtube.com/watch?v={video_id}", download=False)
    subs = info.get("subtitles") or {}
    auto = info.get("automatic_captions") or {}
    track = lang = None
    for L in LANG_PREF:
        if L in subs:
            track, lang = subs[L], L; break
    if track is None:
        for L in LANG_PREF:
            if L in auto:
                track, lang = auto[L], L; break
    if track is None:
        for src in (subs, auto):
            for L, tr in src.items():
                track, lang = tr, L; break
            if track:
                break
    if not track:
        return None, None
    fmt = next((f for f in track if f.get("ext") == "json3"), track[0])
    req = urllib.request.Request(fmt["url"], headers={"User-Agent": "Mozilla/5.0"})
    raw = urllib.request.urlopen(req, timeout=25).read().decode("utf-8", "ignore")
    lines = []
    try:
        j = json.loads(raw)
        for ev in j.get("events", []):
            seg = "".join(s.get("utf8", "") for s in ev.get("segs", []) or [])
            seg = seg.replace("\n", " ").strip()
            if seg:
                lines.append(seg)
    except json.JSONDecodeError:
        raw = re.sub(r"<[^>]+>", "", raw)
        for ln in raw.splitlines():
            ln = ln.strip()
            if ln and "-->" not in ln and not ln.isdigit() and ln != "WEBVTT":
                lines.append(ln)
    return ("\n".join(lines).strip() or None), lang


def safe_filename(name, maxlen=60):
    name = re.sub(r'[\\/:*?"<>|\n\r\t]', "_", name).strip()
    return (name[:maxlen]).strip() or "untitled"


def main():
    # 출력 폴더 결정: 인자로 경로를 주면 그곳, 아니면 바탕화면/유튜브_자막_모음
    if len(sys.argv) > 1:
        out_dir = Path(sys.argv[1]).expanduser()
    else:
        out_dir = find_desktop() / FOLDER_NAME
    txt_dir = out_dir / "텍스트"
    txt_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / f"{FOLDER_NAME}.docx"
    csv_path = out_dir / "결과요약.csv"

    log(f"저장 위치: {out_dir}\n")

    # 대상 목록 만들기 (재생목록 펼치기 + 중복 제거, 순서 유지)
    video_ids = []
    seen = set()

    def add(v):
        if v and v not in seen:
            seen.add(v); video_ids.append(v)

    for v in VIDEO_IDS:
        add(v)
    if EXPAND_PLAYLISTS:
        for p in PLAYLISTS:
            for v in expand_playlist(p):
                add(v)

    total = len(video_ids)
    log(f"총 {total}개 영상 처리 시작\n")

    # 워드 문서 준비
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        log("[오류] python-docx 미설치 → pip install python-docx")
        sys.exit(1)
    doc = Document()
    doc.add_heading("유튜브 자막 모음", level=0)
    doc.add_paragraph(f"총 {total}개 영상 · 한국어 자막 우선(없으면 영어)").italic = True
    doc.add_paragraph("")

    results = []
    for i, vid in enumerate(video_ids, 1):
        title, author = get_title(vid)
        text = lang = method = None
        for attempt in range(RETRY + 1):
            text, lang, method = fetch_transcript(vid)
            if text:
                break
            time.sleep(1.5 * (attempt + 1))
        ok = bool(text)
        nchars = len(text) if ok else 0
        results.append((i, vid, title, author, ok, method, lang, nchars))

        # 영상별 텍스트 파일
        fname = f"{i:04d}_{safe_filename(title)}.txt"
        with open(txt_dir / fname, "w", encoding="utf-8") as f:
            f.write(f"제목: {title}\n채널: {author}\n주소: https://www.youtube.com/watch?v={vid}\n")
            f.write(f"언어: {lang or '-'} · 방법: {method}\n{'='*50}\n\n")
            f.write(text if ok else "[자막을 가져오지 못함: 자막 비공개이거나 요청 차단]")

        # 워드 섹션
        doc.add_heading(f"{i}. {title}", level=1)
        meta = doc.add_paragraph()
        meta.add_run(f"https://www.youtube.com/watch?v={vid}").font.size = Pt(9)
        if author:
            meta.add_run(f"   ·  {author}").font.size = Pt(9)
        if ok:
            for para in text.split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        else:
            r = doc.add_paragraph().add_run("⚠ 자막을 가져오지 못했습니다 (자막 비공개 또는 요청 차단).")
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        doc.add_page_break()

        log(f"[{i:>3}/{total}] {'OK ' if ok else '실패'} {vid} · {lang or '-':5} · {nchars:>6}자 · {title[:38]}")
        time.sleep(SLEEP_SEC)

    doc.save(str(docx_path))

    ok_n = sum(1 for r in results if r[4])
    with open(csv_path, "w", encoding="utf-8-sig") as f:
        f.write("번호,영상ID,제목,채널,성공,방법,언어,글자수\n")
        for (i, vid, title, author, ok, method, lang, n) in results:
            f.write(f"{i},{vid},{title.replace(',', ' ')},{author.replace(',', ' ')},"
                    f"{'성공' if ok else '실패'},{method},{lang or ''},{n}\n")

    log("\n" + "=" * 52)
    log(f"완료: 성공 {ok_n} / 전체 {total}")
    log(f"📄 워드 파일 : {docx_path}")
    log(f"📁 영상별 txt: {txt_dir}")
    log(f"📊 요약 CSV  : {csv_path}")
    if ok_n < total:
        log(f"\n실패 {total-ok_n}개 = 자막이 꺼진 영상이거나 일시 차단. 결과요약.csv에서 확인하고,")
        log("잠시 뒤 다시 실행하면 일시 차단분은 대개 잡힙니다.")


if __name__ == "__main__":
    main()
